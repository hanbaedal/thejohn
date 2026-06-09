# -*- coding: utf-8 -*-
"""더존(thejohn) 이용 서버·인프라 설명 Word 문서 생성."""
from __future__ import annotations

import os
import sys
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from file_inventory_data import build_file_sections  # noqa: E402

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, "docs")
OUT = os.path.join(DOCS, "thejohn-server-infrastructure.docx")
TODAY = date.today().strftime("%Y-%m-%d")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_table(doc, headers, rows, small=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        if small:
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            if small:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
    doc.add_paragraph()


def append_file_inventory(doc):
    sections = build_file_sections(ROOT)
    total = sum(len(items) for _, items in sections)

    doc.add_page_break()
    add_heading(doc, "10. 프로그램 파일 목록", 1)
    note = doc.add_paragraph()
    note.add_run(
        "※ 아래 표는 본 문서 맨 끝(10장)에 이어집니다. "
        "브라우저 「보기」에서는 뒤쪽이 잘 안 보일 수 있으니 "
        "「저장」 후 Word로 열거나, 문서 다운로드의 「프로그램 파일 목록」을 이용하세요."
    ).bold = True
    doc.add_paragraph(
        f"저장소에 포함된 프로그램·문서·자산 파일을 구분별로 정리한 목록입니다. "
        f"(node_modules, .git, server/public 등 빌드·시스템 폴더 제외, 총 {total}개)"
    )

    global_no = 1
    for section_title, items in sections:
        add_heading(doc, section_title, 2)
        rows = []
        for rel, purpose in items:
            rows.append([str(global_no), rel, purpose])
            global_no += 1
        add_table(doc, ["순번", "파일명", "용도"], rows, small=True)


def build_word(path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("더존(thejohn) 업무 시스템\n이용 서버·인프라 설명")
    run.bold = True
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"작성 기준일: {TODAY}\n관리·운영 담당자용").font.size = Pt(11)
    doc.add_page_break()

    add_heading(doc, "1. 문서 목적", 1)
    doc.add_paragraph(
        "본 문서는 더존(thejohn) 홈페이지·업무 시스템이 실제로 어떤 서버와 "
        "외부 서비스를 이용해 동작하는지, 비개발 담당자도 이해할 수 있도록 정리한 것입니다."
    )
    doc.add_paragraph(
        "화면·메뉴·역할별 사용법은 사용자 매뉴얼, 코드·디렉터리 상세는 기술 구조 문서를 참고하세요."
    )

    add_heading(doc, "2. 한 줄 요약", 1)
    doc.add_paragraph(
        "Render(클라우드)에서 Node.js 서버가 웹 화면과 API를 제공하고, "
        "데이터는 MongoDB Atlas(클라우드 DB)에 저장됩니다. "
        "접속 주소는 thejohn.co.kr 도메인을 사용합니다."
    )

    add_heading(doc, "3. 이용 서버·서비스 목록", 1)
    add_table(
        doc,
        ["구분", "서비스", "역할"],
        [
            ["웹·API 서버", "Render", "HTML/JS 화면 제공, 로그인·주문·상품 등 API 처리"],
            ["데이터베이스", "MongoDB Atlas", "직원·거래처·상품·주문·접속 로그 등 저장"],
            ["도메인·DNS", "가비아", "thejohn.co.kr 주소를 Render 서버로 연결"],
            ["소스·배포", "GitHub", "프로그램 소스 저장, main 푸시 시 Render 자동 배포"],
            ["SMS(선택)", "SOLAPI", "주문 알림 등 문자 발송(설정된 경우)"],
            ["로컬 개발", "본인 PC", "개발·테스트 시 localhost:3000 에서 실행"],
        ],
    )

    add_heading(doc, "4. 웹·API 서버 (Render)", 1)
    doc.add_paragraph(
        "운영 중인 프로그램은 Render라는 클라우드 호스팅에서 Node.js(Express)로 실행됩니다. "
        "브라우저에 보이는 페이지(HTML, JS, CSS)와 백엔드 API(/api/...)가 "
        "같은 서버 한 대에서 함께 제공됩니다."
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["접속 주소", "https://thejohn.co.kr , https://www.thejohn.co.kr"],
            ["Render 주소", "https://thejohn.onrender.com (호스팅 기본 URL)"],
            ["실행 방식", "server 폴더에서 npm start (빌드 후 정적 파일 포함)"],
            ["상태 확인", "GET /api/health (서비스·DB 연결 점검)"],
            ["배포", "GitHub main 브랜치 푸시 → Render 자동 반영"],
        ],
    )
    doc.add_paragraph("요청 흐름 (간단):")
    for line in [
        "① 사용자가 thejohn.co.kr/어떤페이지.html 접속",
        "② Render 서버가 HTML·JS·CSS 전달",
        "③ 페이지가 같은 서버의 /api/* 로 데이터 요청 (로그인 토큰 포함)",
        "④ 서버가 MongoDB에서 조회·저장 후 결과 반환",
        "⑤ PDF(발주서·거래명세)는 서버에서 생성해 보기·저장",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "5. 데이터베이스 (MongoDB Atlas)", 1)
    doc.add_paragraph(
        "업무 데이터는 MongoDB Atlas라는 클라우드 데이터베이스에 저장됩니다. "
        "Render 서버만 Atlas에 접속할 수 있도록 연결 문자열(MONGODB_URI)이 설정되어 있습니다."
    )
    add_table(
        doc,
        ["컬렉션(예)", "저장 내용"],
        [
            ["staff", "내부 직원(슈퍼바이저·관리자) 계정"],
            ["vendors", "거래처(업체) 로그인·등급·담당 관리자"],
            ["products", "상품 마스터·사진"],
            ["orders", "주문·발주 정보"],
            ["access_logs", "접속·페이지 방문·세션 통계"],
            ["transaction_manual", "수기 거래명세서"],
        ],
    )

    add_heading(doc, "6. 도메인 (thejohn.co.kr)", 1)
    doc.add_paragraph(
        "사용자가 입력하는 주소 thejohn.co.kr 은 가비아에서 관리하는 도메인입니다. "
        "DNS 설정을 통해 실제 처리는 Render 웹 서버가 담당합니다. "
        "따라서 브라우저 주소창에는 thejohn.co.kr 이 보이지만, "
        "프로그램 실행·API는 Render에서 이루어집니다."
    )

    add_heading(doc, "7. GitHub와 배포", 1)
    doc.add_paragraph(
        "프로그램 소스는 GitHub 저장소(hanbaedal/thejohn)에 있습니다. "
        "수정 사항을 main 브랜치에 반영(푸시)하면 Render가 새 버전을 받아 "
        "빌드·재시작하여 운영 서버에 적용합니다. 배포에는 보통 1~3분 정도 걸릴 수 있습니다."
    )

    add_heading(doc, "8. 로컬 개발 환경", 1)
    doc.add_paragraph(
        "개발자 PC에서는 server 폴더에서 npm start 로 http://localhost:3000 에서 "
        "동일한 구조로 테스트합니다. HTML 파일을 file:// 로 직접 열면 API가 동작하지 않습니다."
    )

    add_heading(doc, "9. 관련 문서", 1)
    add_table(
        doc,
        ["문서", "파일", "용도"],
        [
            ["사용자 매뉴얼", "thejohn-user-manual.docx", "화면·역할별 조작 방법"],
            ["시스템 구조 설명서", "thejohn-system-structure-management.docx", "관리·기획용 구조"],
            ["시스템 구조 요약", "thejohn-system-structure-management.pptx", "회의·보고용 슬라이드"],
            ["기술 구조", "thejohn-system-technical.docx", "개발·운영 상세(ARCHITECTURE.md 기반)"],
            ["이용 서버·인프라", "thejohn-server-infrastructure.docx", "본 문서"],
            ["프로그램 파일 목록", "thejohn-file-inventory.docx", "순번·파일명·용도 전체 목록"],
        ],
    )
    doc.add_paragraph(
        "웹에서 받기: 로그인(관리자·슈퍼바이저) → 업무관리 → 문서 다운로드"
    )

    append_file_inventory(doc)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        f"문서 버전: {TODAY} 기준. 상세 설정은 저장소 docs/ARCHITECTURE.md, deploy/ 폴더를 참고하세요."
    ).font.size = Pt(9)

    doc.save(path)
    print("Server infrastructure doc:", path)


def main():
    os.makedirs(DOCS, exist_ok=True)
    build_word(OUT)


if __name__ == "__main__":
    main()
