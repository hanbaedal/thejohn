# -*- coding: utf-8 -*-
"""더존(thejohn) 사용자 매뉴얼 Word(.docx) 생성."""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, "docs")
TODAY = date.today().strftime("%Y-%m-%d")
OUT = os.path.join(DOCS, "thejohn-user-manual.docx")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, steps):
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")


def build_user_manual(path):
    doc = Document()
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], margin, Cm(2.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("더존(thejohn) 웹 시스템\n사용자 매뉴얼")
    r.bold = True
    r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"작성 기준일: {TODAY}\n버전: 1.1").font.size = Pt(11)
    doc.add_page_break()

    add_heading(doc, "목차 안내", 1)
    doc.add_paragraph(
        "본 매뉴얼은 로그인부터 주문·발주·거래명세·상품·업체 관리까지 "
        "화면에서 하는 작업을 역할별로 설명합니다. "
        "시스템 구조·서버 설명은 별도 문서(시스템 구조 설명서)를 참고하세요."
    )
    add_table(
        doc,
        ["장", "제목"],
        [
            ["1", "시작하기 (접속·로그인)"],
            ["2", "이용자 종류와 메뉴"],
            ["3", "게스트(비로그인) 이용"],
            ["4", "거래처(업체) 이용"],
            ["5", "관리자·슈퍼바이저 공통"],
            ["6", "홈페이지 관리"],
            ["7", "상품 관리"],
            ["8", "업체 관리"],
            ["9", "주문서 관리"],
            ["10", "수기 거래명세서"],
            ["11", "PDF 보기와 저장"],
            ["12", "업무관리(슈퍼바이저)"],
            ["13", "자주 묻는 질문"],
        ],
    )

    add_heading(doc, "1. 시작하기", 1)
    add_heading(doc, "1.1 접속", 2)
    add_bullets(
        doc,
        [
            "권장 주소: https://www.thejohn.co.kr (또는 thejohn.co.kr)",
            "PC·태블릿·스마트폰의 Chrome, Edge, Safari 등 최신 브라우저 사용",
            "HTML 파일을 PC에 저장해 더블클릭(file://)으로 열면 로그인·주문이 되지 않습니다. 반드시 위 주소로 접속하세요.",
        ],
    )
    add_heading(doc, "1.2 로그인", 2)
    add_steps(
        doc,
        [
            "상단 또는 안내에 따라 login.html(로그인) 화면으로 이동합니다.",
            "아이디·비밀번호를 입력하고 로그인합니다.",
            "로그인 성공 후 역할에 따라 메인(홈) 또는 그룹 마케팅 관리(work-hub)로 이동합니다.",
            "우측 상단 로그아웃 버튼으로 종료할 수 있습니다.",
        ],
    )
    doc.add_paragraph(
        "※ 아이디·비밀번호 분실 시 슈퍼바이저 또는 담당 관리자에게 문의하세요. "
        "초기 관리자 비밀번호는 서버 환경 설정으로만 변경됩니다."
    )

    add_heading(doc, "2. 이용자 종류와 메뉴", 1)
    add_table(
        doc,
        ["역할", "누가 쓰나", "로그인 후 주요 메뉴"],
        [
            ["슈퍼바이저", "본사 총괄", "그룹 마케팅 관리 6개(홈·홈관리·상품·업체·주문·업무)"],
            [
                "관리자",
                "영업·담당 직원",
                "주문 권한 있음: 5개 / 없음: 4개(주문서 관리 제외)",
            ],
            ["거래처", "등록된 업체", "상품 목록·장바구니·주문"],
            ["게스트", "로그인 없이 열람", "상품 목록(가격 미표시)"],
        ],
    )
    add_heading(doc, "2.1 그룹 마케팅 관리(work-hub)", 2)
    doc.add_paragraph(
        "관리자·슈퍼바이저가 로그인하면 「그룹 마케팅 관리」 화면에서 아래 메뉴로 이동합니다."
    )
    add_table(
        doc,
        ["메뉴", "설명"],
        [
            ["홈페이지", "고객용 메인·사업부문 소개"],
            ["홈페이지 관리하기", "소식·게시판·문의 답변"],
            ["상품관리", "상품 등록·수정·목록"],
            ["업체관리", "거래처 등록·수정·예비거래처 등"],
            ["주문서 관리", "발주·거래명세·수기 명세 (권한 있는 관리자만)"],
            ["업무관리", "직원·통계·문서 다운로드 (슈퍼바이저)"],
        ],
    )

    add_heading(doc, "3. 게스트(비로그인) 이용", 1)
    add_bullets(
        doc,
        [
            "메인·사업부문 페이지·상품 목록을 볼 수 있습니다.",
            "가격은 표시되지 않을 수 있습니다.",
            "주문·관리 메뉴는 사용할 수 없습니다.",
        ],
    )

    add_heading(doc, "4. 거래처(업체) 이용", 1)
    add_heading(doc, "4.1 상품 보기·주문", 2)
    add_steps(
        doc,
        [
            "업체 계정으로 로그인합니다.",
            "상품(products) 메뉴에서 담당 관리자가 등록한 상품을 확인합니다. 타 관리자 상품은 제한될 수 있습니다.",
            "수량을 선택해 장바구니에 담습니다.",
            "주문하기(장바구니)에서 품목·수량을 확인하고 주문을 제출합니다.",
            "주문 완료 후 발주서 PDF 저장 등 안내가 표시될 수 있습니다.",
        ],
    )
    add_heading(doc, "4.2 주문 이력·PDF", 2)
    add_bullets(
        doc,
        [
            "주문 이력 화면에서 과거 주문을 확인합니다.",
            "「PDF 저장」을 누르면 발주서 파일이 PC에 저장됩니다.",
            "담당 관리자·주문 권한이 없으면 주문 버튼이 비활성화될 수 있습니다.",
        ],
    )

    add_heading(doc, "5. 관리자·슈퍼바이저 공통", 1)
    add_bullets(
        doc,
        [
            "헤더 날짜·로그아웃: 모든 관리 화면 우측 상단",
            "← 그룹 마케팅 관리: 하위 메뉴에서 상위 허브로 돌아가기",
            "목록 화면: 기간·담당 관리자 필터 후 「조회」 버튼 사용",
            "항목 클릭 시 상세 모달에서 품목·합계 확인",
        ],
    )

    add_heading(doc, "6. 홈페이지 관리", 1)
    doc.add_paragraph("경로: 그룹 마케팅 관리 → 홈페이지 관리하기")
    add_table(
        doc,
        ["기능", "하는 일"],
        [
            ["최근소식 입력", "고객센터 소식 등록·수정"],
            ["자유게시판", "게시글 관리"],
            ["문의사항 답변", "고객 문의 조회·답변 등록"],
        ],
    )

    add_heading(doc, "7. 상품 관리", 1)
    add_steps(
        doc,
        [
            "그룹 마케팅 관리 → 상품관리",
            "상품 등록: 품목코드, 품명, 규격, 가격, 사업부문 등 입력 후 저장",
            "상품 사진: 1장만 등록 (540×540 JPEG, 1MB 이하). 앨범·카메라로 선택",
            "규격 옆 「상품정보」: 식품 표시사항 등 항목별 입력·저장 (product_info)",
            "상품 리스트: 검색·수정·삭제",
            "홈페이지 사업부문: 상품 클릭 시 같은 부문 상품 상세를 위·아래로 스크롤하며 볼 수 있음",
        ],
    )
    doc.add_paragraph(
        "※ 슈퍼바이저가 수기 거래명세 작성 시 「공급자」를 선택하면, "
        "해당 관리자가 등록한 상품만 품목 선택 목록에 나타납니다."
    )

    add_heading(doc, "8. 업체 관리", 1)
    add_steps(
        doc,
        [
            "그룹 마케팅 관리 → 업체관리",
            "업체 등록: 로그인 아이디, 업체명, 담당 관리자, 등급·연락처 등",
            "업체 리스트: 수정·상세 보기",
            "예비거래처·엑셀 일괄 등록(권한·메뉴 있는 경우): 안내에 따라 업로드",
            "업체 메일 발송(메뉴 있는 경우): 수신 업체·첨부 선택 후 발송",
        ],
    )

    add_heading(doc, "9. 주문서 관리", 1)
    doc.add_paragraph("경로: 그룹 마케팅 관리 → 주문서 관리")
    add_table(
        doc,
        ["메뉴", "용도"],
        [
            ["발주서 리스트", "기간·담당별 발주 조회·상세·PDF"],
            ["발주서 PDF로 보기", "목록에서 PDF 보기(모달)·PDF 저장"],
            ["거래명세서 PDF로 보기", "주문 건별 거래명세 PDF·인쇄"],
            ["거래명세서 수기 작성", "주문 없이 명세 작성"],
            ["수기 거래명세서 목록", "저장 건 조회·수정·PDF 보기"],
        ],
    )
    add_heading(doc, "9.1 발주 목록·상세", 2)
    add_steps(
        doc,
        [
            "시작일·종료일(및 슈퍼바이저: 담당 관리자) 선택 후 조회",
            "목록에서 주문 번호·업체명·금액 확인",
            "행 클릭 또는 PDF 보기: 상세·문서 확인",
            "관리자 주문 목록: PDF 보기(모달), 상세의 PDF 저장",
        ],
    )
    add_heading(doc, "9.2 PDF 목록 화면", 2)
    add_bullets(
        doc,
        [
            "「PDF 보기」: 화면 가운데 모달로 문서 확인 (자동 다운로드 없음)",
            "「PDF 저장」: PC에 파일로 저장",
            "거래명세 화면: 「인쇄」는 인쇄용 새 창(브라우저 인쇄 대화상자)",
        ],
    )

    add_heading(doc, "10. 수기 거래명세서", 1)
    add_heading(doc, "10.1 작성", 2)
    add_steps(
        doc,
        [
            "주문서 관리 → 거래명세서 수기 작성",
            "발행일: 년·월·일 선택",
            "슈퍼바이저: 공급자(발행 관리자) 선택 — 인감·상품 범위 결정",
            "관리자: 공급자는 본인 고정(화면에 표시만 될 수 있음)",
            "「업체 선택」으로 등록 업체 불러오기, 품목은 품목코드/품명 칸 클릭 후 상품 선택",
            "수량: − / 숫자 / + 버튼 또는 직접 입력",
            "메모: 목록에 표시할 제목(선택), 비고: 내부 메모(PDF에 안 나갈 수 있음)",
            "「저장」: DB 저장 후 PDF 저장 버튼 활성화",
            "「PDF 미리보기」: 저장 전 모달로 확인",
            "「PDF 저장」: 저장된 건만 파일 다운로드",
        ],
    )
    add_heading(doc, "10.2 목록", 2)
    add_steps(
        doc,
        [
            "주문서 관리 → 수기 거래명세서 목록",
            "슈퍼바이저: 발행 관리자 필터 가능 / 관리자: 본인 작성분만",
            "「PDF 보기」: 모달 / 「수정」: 작성 화면으로 이동",
        ],
    )

    add_heading(doc, "11. PDF 보기와 저장", 1)
    add_table(
        doc,
        ["버튼", "동작"],
        [
            ["PDF 보기", "모달 창에서 확인. ×, 바깥 클릭, Esc로 닫기"],
            ["PDF 저장", "거래명세서_업체명_날짜.pdf 형식 등으로 PC 저장"],
            ["PDF 미리보기", "수기 작성 중 저장 전 확인(모달)"],
        ],
    )
    doc.add_paragraph(
        "모달이 안 열리면 브라우저 팝업 차단이 아닌 경우가 많습니다. "
        "그래도 안 되면 다른 브라우저로 시도하거나 PDF 저장으로 확인하세요."
    )

    add_heading(doc, "12. 업무관리(슈퍼바이저)", 1)
    doc.add_paragraph("경로: 그룹 마케팅 관리 → 업무관리 (관리자관리 허브)")
    doc.add_paragraph(
        "메뉴는 아래와 같이 3그룹으로 구성됩니다. (관리자 등록/리스트 → 통계 2종 → SOLAPI·문서)"
    )
    add_heading(doc, "12.1 관리자 계정", 2)
    add_table(
        doc,
        ["메뉴", "용도"],
        [
            ["관리자 등록", "내부 계정 생성·주문 권한·인감 등 설정"],
            ["관리자 리스트", "등록된 관리자 조회·수정"],
        ],
    )
    add_heading(doc, "12.2 통계", 2)
    add_table(
        doc,
        ["메뉴", "용도"],
        [
            [
                "접속·이용 통계",
                "관리자·등록 업체·게스트의 로그인, 페이지 방문, 세션(머무른 시간). 기간 필터",
            ],
            [
                "디비사용 통계",
                "담당 관리자별 DB 저장 용량, 이용 시간(접속 로그), 상품·업체·발주 건수. 기간 필터",
            ],
        ],
    )
    add_heading(doc, "12.3 알림·문서", 2)
    add_table(
        doc,
        ["메뉴", "용도"],
        [
            [
                "SOLAPI 이용 현황",
                "업체 주문 시 관리자 SMS(SOLAPI) 발송 횟수. 담당 관리자·등록 업체별 집계. 기능 적용 이후 기록만 포함",
            ],
            [
                "문서 다운로드",
                "사용자 매뉴얼·시스템 구조 설명서(Word/PPT)·기술 구조(Markdown) 받기",
            ],
        ],
    )
    doc.add_paragraph(
        "※ 발주서·거래명세는 「주문서 관리」 메뉴에서 이용합니다. 업무관리와 별도입니다."
    )

    add_heading(doc, "13. 자주 묻는 질문", 1)
    faq = [
        ("로그인이 안 됩니다.", "아이디·비밀번호 확인, Caps Lock, 슈퍼바이저에 계정 상태 문의."),
        ("주문서 관리 메뉴가 없습니다.", "관리자는 슈퍼바이저가 「주문 권한」을 켜야 합니다."),
        ("업체가 상품을 주문할 수 없습니다.", "업체의 담당 관리자·주문 권한·상품 노출 여부 확인."),
        ("PDF가 다운로드만 됩니다.", "「PDF 보기」 버튼 사용. 「PDF 저장」과 다릅니다."),
        ("수기 명세 PDF 저장이 비활성입니다.", "먼저 「저장」으로 문서를 저장한 뒤 PDF 저장."),
        ("화면이 깨지거나 API 오류.", "주소가 www.thejohn.co.kr 인지 확인, 로그아웃 후 재로그인."),
        (
            "SOLAPI 이용 현황에 기록이 없습니다.",
            "기능 적용 이후 주문 알림이 발생한 건만 집계됩니다. 기간을 넓혀 조회해 보세요.",
        ),
        (
            "문서 다운로드 내용이 화면과 다릅니다.",
            "Word/PPT는 스크립트로 생성한 정적 파일입니다. 최신 반영 후 재배포가 필요할 수 있습니다.",
        ),
    ]
    for q, a in faq:
        p = doc.add_paragraph()
        p.add_run(q + " ").bold = True
        p.add_run(a)

    add_heading(doc, "부록: 문의", 1)
    doc.add_paragraph(
        "시스템 장애·계정·권한: 내부 슈퍼바이저 또는 IT 담당자\n"
        "운영 URL·배포: thejohn.co.kr / Render 호스팅\n"
        "본 매뉴얼 파일: docs/thejohn-user-manual.docx"
    )

    doc.save(path)
    print("User manual:", path)


def main():
    os.makedirs(DOCS, exist_ok=True)
    build_user_manual(OUT)


if __name__ == "__main__":
    main()
