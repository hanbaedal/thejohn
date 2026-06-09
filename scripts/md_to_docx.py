# -*- coding: utf-8 -*-
"""Markdown → Word(docx) 변환 공통."""
from __future__ import annotations

import re
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

TODAY = date.today().strftime("%Y-%m-%d")


def parse_table_row(line):
    line = line.strip()
    if not line.startswith("|"):
        return None
    parts = [p.strip() for p in line.strip("|").split("|")]
    if all(re.match(r"^[-:]+$", p) for p in parts):
        return "separator"
    return parts


def build_from_markdown(md_path, out_path, doc_title=None, source_note=None):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], margin, Cm(2.5))

    if doc_title:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run(doc_title)
        r.bold = True
        r.font.size = Pt(20)
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note = f"작성 기준일: {TODAY}"
        if source_note:
            note += "\n" + source_note
        sub.add_run(note).font.size = Pt(10)
        doc.add_page_break()

    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        headers = table_buf[0]
        rows = table_buf[1:]
        if not headers:
            table_buf = []
            return
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        for ci, h in enumerate(headers):
            tbl.rows[0].cells[ci].text = h
        for ri, row in enumerate(rows):
            for ci in range(len(headers)):
                val = row[ci] if ci < len(row) else ""
                tbl.rows[ri + 1].cells[ci].text = val
        doc.add_paragraph()
        table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|"):
            row = parse_table_row(line)
            if row == "separator":
                i += 1
                continue
            if row:
                table_buf.append(row)
            i += 1
            continue
        else:
            flush_table()

        if stripped in ("---", "***"):
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    flush_table()
    flush_code()
    doc.save(out_path)
