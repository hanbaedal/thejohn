# -*- coding: utf-8 -*-
"""프로그램 파일 목록 Word 문서 (thejohn-file-inventory.docx)."""
from __future__ import annotations

import os
import sys
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from file_inventory_data import add_file_inventory_table, build_file_sections  # noqa: E402

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, "docs")
OUT = os.path.join(DOCS, "thejohn-file-inventory.docx")
TODAY = date.today().strftime("%Y-%m-%d")


def build_word(path):
    sections = build_file_sections(ROOT)
    total = sum(len(items) for _, items in sections)

    doc = Document()
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], margin, Cm(2.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("더존(thejohn) 프로그램 파일 목록")
    run.bold = True
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"작성 기준일: {TODAY}\n총 {total}개 파일").font.size = Pt(11)

    doc.add_paragraph(
        "저장소에 포함된 화면·스크립트·서버·문서·이미지 파일을 구분별로 정리했습니다. "
        "(node_modules, .git, server/public 제외)"
    )

    global_no = 1
    for section_title, items in sections:
        doc.add_heading(section_title, level=1)
        rows = []
        for rel, purpose in items:
            rows.append([str(global_no), rel, purpose])
            global_no += 1
        add_file_inventory_table(doc, rows)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(f"문서 버전: {TODAY} 기준.").font.size = Pt(9)

    doc.save(path)
    print("File inventory doc:", path)


def main():
    os.makedirs(DOCS, exist_ok=True)
    build_word(OUT)


if __name__ == "__main__":
    main()
