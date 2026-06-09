# -*- coding: utf-8 -*-
"""저장소 파일 목록·용도 — server-infrastructure docx 부록용."""
from __future__ import annotations

import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

SKIP_DIRS = {
    ".git",
    ".cursor",
    "node_modules",
    "server/public",
    "__pycache__",
}

SKIP_FILES = {
    ".env",
    "package-lock.json",
}

SKIP_SUFFIXES = {".pyc", ".pyo"}

PURPOSE_MAP = {
    "index.html": "메인 홈페이지",
    "login.html": "로그인 화면",
    "company.html": "회사소개 메인",
    "company-driedfish.html": "회사소개 — 건어물",
    "company-drink.html": "회사소개 — 음료수",
    "company-frozen.html": "회사소개 — 냉동식품",
    "company-grocery.html": "회사소개 — 공산품",
    "company-jeongyuk.html": "회사소개 — 정육",
    "company-seafood.html": "회사소개 — 냉동수산물",
    "products.html": "사업부문(상품 카탈로그)",
    "product-detail.html": "상품 상세",
    "product-manage.html": "상품관리 허브",
    "product-register.html": "상품 내용 등록",
    "product-edit.html": "상품 등록 수정",
    "product-list-admin.html": "상품 리스트(관리)",
    "product-new-list.html": "신규상품 리스트",
    "product-new-register.html": "신규상품 등록",
    "vendor-manage.html": "업체관리 허브",
    "vendor-register.html": "업체 등록",
    "vendor-edit.html": "업체 수정",
    "vendor-detail.html": "업체 상세",
    "vendor-list-admin.html": "업체 리스트(관리)",
    "vendor-new-list.html": "신규업체 리스트",
    "vendor-new-register.html": "신규업체 등록",
    "vendor-prospect-list.html": "예비업체 리스트",
    "vendor-prospect-finder.html": "예비 업체 찾기",
    "vendor-excel-import.html": "예비업체 엑셀 자료찾기",
    "vendor-email-broadcast.html": "업체별 이메일 발송",
    "vendor-email-history.html": "이메일 발송 내역",
    "vendor-dm-print.html": "업체별 DM 출력",
    "cart.html": "주문서(장바구니) 보기",
    "order-manage-hub.html": "주문서 관리 허브",
    "order-list-admin.html": "주문서 관리 목록",
    "supervisor-order-list.html": "발주서 관리(슈퍼바이저)",
    "supervisor-order-pdf.html": "발주서 PDF 보기",
    "supervisor-transaction-list.html": "거래명세서 목록",
    "supervisor-transaction-pdf.html": "거래명세서 PDF 보기",
    "transaction-manual-register.html": "수기 거래명세서 작성",
    "transaction-manual-list.html": "수기 거래명세서 목록",
    "work-hub.html": "그룹 마케팅 관리 허브",
    "staff-manage-hub.html": "관리자관리 허브",
    "staff-manage.html": "관리자 등록",
    "staff-list-admin.html": "관리자 리스트",
    "staff-self-edit.html": "관리자 정보 수정",
    "homepage-manage-hub.html": "홈페이지 관리 허브",
    "support.html": "고객센터 메인",
    "support-library.html": "최근소식(고객)",
    "support-news-admin.html": "최근소식 입력(관리)",
    "support-qna.html": "자유게시판(고객)",
    "support-qna-admin.html": "자유게시판 관리",
    "support-inquiry.html": "문의사항(고객)",
    "support-inquiry-reply.html": "문의사항 답변(관리)",
    "support-partners.html": "파트너회사 소개",
    "support-partner-detail.html": "파트너 업체 상세",
    "supervisor-access-stats.html": "접속통계 관리",
    "supervisor-usage-stats.html": "접속·이용 통계",
    "supervisor-db-stats.html": "DB 사용 통계",
    "supervisor-solapi-stats.html": "SOLAPI 이용 현황",
    "system-structure-docs.html": "문서 다운로드",
    "thejhon-api.js": "REST API 클라이언트",
    "api-config.js": "API 베이스 URL 설정",
    "auth.js": "로그인·역할·페이지 가드",
    "auth-storage.js": "토큰·세션 저장소",
    "oauth-config.js": "OAuth 설정",
    "nav.js": "공통 네비게이션",
    "admin-header.js": "관리자 헤더·메뉴",
    "site-brand-boot.js": "로고·파비콘 브랜드 부트",
    "footer-company.js": "푸터 회사정보",
    "footer-social.js": "푸터 SNS 링크",
    "home-intro-media.js": "홈 인트로 영상·음악",
    "company-greeting.js": "회사소개 인사말",
    "login.js": "로그인 폼 처리",
    "products.js": "사업부문 페이지 로직",
    "product-catalog.js": "상품 카탈로그 공통",
    "product-detail.js": "상품 상세 페이지",
    "product-register.js": "상품 등록 페이지",
    "product-edit.js": "상품 수정 페이지",
    "product-form-shared.js": "상품 폼 공통 로직",
    "product-list-admin.js": "상품 리스트(관리)",
    "product-new-list.js": "신규상품 리스트",
    "product-new-register.js": "신규상품 등록",
    "product-cover-cache.js": "상품 커버 이미지 캐시",
    "product-info-modal.js": "상품정보(식품표시) 모달",
    "vendor-register.js": "업체 등록",
    "vendor-edit.js": "업체 수정",
    "vendor-detail.js": "업체 상세",
    "vendor-form-shared.js": "업체 폼 공통",
    "vendor-admin-shared.js": "업체 관리 공통",
    "vendor-list-admin.js": "업체 리스트(관리)",
    "vendor-new-list.js": "신규업체 리스트",
    "vendor-new-register.js": "신규업체 등록",
    "vendor-prospect-list.js": "예비업체 리스트",
    "vendor-prospect-finder.js": "예비업체 검색",
    "vendor-prospect-picker.js": "예비업체 선택 UI",
    "vendor-excel-import.js": "엑셀 예비업체 가져오기",
    "vendor-excel-import-map.js": "엑셀 열→필드 매핑",
    "vendor-email-broadcast.js": "업체 이메일 일괄 발송",
    "vendor-email-history.js": "이메일 발송 내역",
    "vendor-dm-print.js": "업체 DM 인쇄",
    "vendor-cart.js": "업체 장바구니",
    "vendor-order-modal.js": "업체 주문 모달",
    "cart.js": "주문서(장바구니) 페이지",
    "catalog-order-ui.js": "카탈로그 주문 UI",
    "qty-stepper.js": "수량 스테퍼 컴포넌트",
    "order-ui.js": "주문 상세·PDF UI",
    "order-detail-modal.js": "주문 상세 모달",
    "order-pdf-client.js": "주문 PDF 클라이언트",
    "order-list-admin.js": "주문서 관리 목록",
    "order-manage-hub.js": "주문서 관리 허브",
    "order-manage-pdf-list.js": "주문 PDF 목록",
    "supervisor-order-list.js": "발주서 관리(슈퍼바이저)",
    "supervisor-transaction-list.js": "거래명세서 목록",
    "transaction-manual-register.js": "수기 거래명세 작성",
    "transaction-manual-list.js": "수기 거래명세 목록",
    "work-hub.js": "업무 허브",
    "staff-manage-hub.js": "관리자관리 허브",
    "staff-manage.js": "관리자 등록",
    "staff-list-admin.js": "관리자 리스트",
    "staff-self-edit.js": "관리자 정보 수정",
    "staff-manage-home-boot.js": "홈페이지관리 네비 부트",
    "staff-product-boot.js": "상품관리 네비 부트",
    "staff-vendor-boot.js": "업체관리 네비 부트",
    "homepage-manage-hub.js": "홈페이지 관리 허브",
    "support-hub.js": "고객센터 허브",
    "support-common.js": "고객지원 공통",
    "support-library.js": "최근소식(고객)",
    "support-news-admin.js": "최근소식 관리",
    "support-news-shared.js": "소식 공통 로직",
    "support-qna.js": "자유게시판(고객)",
    "support-qna-admin.js": "자유게시판 관리",
    "support-inquiry.js": "문의사항(고객)",
    "support-inquiry-reply.js": "문의 답변(관리)",
    "support-partners.js": "파트너회사 소개",
    "support-partner-detail.js": "파트너 업체 상세",
    "supervisor-access-stats.js": "접속통계 화면",
    "supervisor-usage-stats.js": "이용 통계 화면",
    "supervisor-db-stats.js": "DB 통계 화면",
    "supervisor-solapi-stats.js": "SOLAPI 통계 화면",
    "address-fields.js": "주소 입력 필드",
    "common.css": "전역 공통 스타일",
    "login.css": "로그인 페이지",
    "home-intro-media.css": "홈 인트로 미디어",
    "admin-manage.css": "관리 화면 공통",
    "staff-manage.css": "관리자 등록·수정",
    "staff-hub.css": "관리자 허브",
    "homepage-manage-hub.css": "홈페이지 관리 허브",
    "work-hub.css": "업무 허브",
    "product-form-page.css": "상품 등록·수정 폼",
    "product-photo-gallery.css": "상품 사진 갤러리",
    "product-info-modal.css": "상품정보 모달",
    "vendor-form-page.css": "업체 등록·수정 폼",
    "vendor-order-modal.css": "업체 주문 모달",
    "vendor-prospect-picker.css": "예비업체 선택",
    "vendor-excel-import.css": "엑셀 가져오기",
    "vendor-dm-print.css": "DM 인쇄 레이아웃",
    "catalog-order-ui.css": "카탈로그 주문 UI",
    "order-manage-layout.css": "주문서 관리 레이아웃",
    "order-list-filters.css": "주문 목록 필터",
    "order-detail-modal.css": "주문 상세 모달",
    "pdf-view-modal.css": "PDF 보기 모달",
    "transaction-manual-register.css": "수기 거래명세 작성",
    "transaction-manual-list.css": "수기 거래명세 목록",
    "support-pages.css": "고객지원 페이지",
    "support-hub.css": "고객센터 허브",
    "support-partners.css": "파트너 소개",
    "address-fields.css": "주소 입력 필드",
    "server/index.js": "Express 앱 진입점",
    "server/db.js": "MongoDB 연결",
    "server/package.json": "서버 의존성·빌드 스크립트",
    "server/routes/auth.js": "로그인·JWT 발급",
    "server/routes/staff.js": "직원(관리자) CRUD",
    "server/routes/products.js": "상품 API",
    "server/routes/vendors.js": "거래처 API",
    "server/routes/vendorProspects.js": "예비 거래처 API",
    "server/routes/vendorNew.js": "신규 업체 등록 API",
    "server/routes/orders.js": "주문·발주·거래명세 PDF",
    "server/routes/transactionManual.js": "수기 거래명세 API",
    "server/routes/vendorEmail.js": "업체 이메일 발송",
    "server/routes/supervisor.js": "슈퍼바이저 통계·집계",
    "server/routes/access.js": "접속·페이지뷰 통계",
    "server/routes/supportNews.js": "최근소식 API",
    "server/routes/supportBoard.js": "자유게시판 API",
    "server/routes/supportInquiry.js": "1:1 문의 API",
    "server/lib/accessLog.js": "접속·세션·이용 통계",
    "server/lib/addressFormat.js": "주소 포맷 정규화",
    "server/lib/image540.js": "상품 이미지 리사이즈·썸네일",
    "server/lib/loginAccount.js": "로그인 계정 공통",
    "server/lib/loginLookup.js": "로그인 ID 조회",
    "server/lib/loginResolve.js": "로그인 계정 해석",
    "server/lib/passwordStore.js": "비밀번호 해시",
    "server/lib/sessionControl.js": "세션·동시 로그인 제어",
    "server/lib/regenerateDocs.js": "문서 재생성 트리거",
    "server/lib/staff.js": "직원 계정·시드",
    "server/lib/staffFields.js": "직원 스키마 필드",
    "server/lib/staffLoginId.js": "직원 로그인 ID",
    "server/lib/staffLoginIdMigration.js": "로그인 ID 마이그레이션",
    "server/lib/staffOrderEnabled.js": "관리자 주문 권한 플래그",
    "server/lib/staffRegisteredBy.js": "직원 등록자 추적",
    "server/lib/staffSealImage.js": "직원 인감 이미지",
    "server/lib/productAccess.js": "상품 접근 권한",
    "server/lib/productDept.js": "상품 사업부문",
    "server/lib/productFields.js": "상품 스키마 필드",
    "server/lib/productInfo.js": "상품정보(식품표시) CRUD",
    "server/lib/vendorAccess.js": "업체 접근 권한",
    "server/lib/vendorCollections.js": "업체 컬렉션 헬퍼",
    "server/lib/vendorExternalLookup.js": "외부 업체 조회",
    "server/lib/vendorFields.js": "업체 스키마 필드",
    "server/lib/vendorNew.js": "신규 업체 등록 로직",
    "server/lib/vendorPricing.js": "업체 등급별 단가",
    "server/lib/vendorProspects.js": "예비업체 도메인 로직",
    "server/lib/vendorProspectImport.js": "예비업체 엑셀 가져오기",
    "server/lib/orderAccess.js": "주문 조회 권한·필터",
    "server/lib/orderDeptLabels.js": "주문 사업부문 라벨",
    "server/lib/orderEnrich.js": "주문·PDF 데이터 보강",
    "server/lib/orderNotify.js": "주문 알림",
    "server/lib/orderPdf.js": "발주서 PDF 생성",
    "server/lib/transactionManual.js": "수기 거래명세 DB·검증",
    "server/lib/transactionIssuer.js": "거래명세 공급자 정보",
    "server/lib/transactionPdf.js": "거래명세서 PDF",
    "server/lib/transactionPdfDouzone.js": "더존 거래명세 PDF 레이아웃",
    "server/lib/solapiSms.js": "SOLAPI SMS 발송",
    "server/lib/solapiLog.js": "SMS 발송 로그",
    "server/lib/supportAuthor.js": "고객지원 작성자 정보",
    "server/lib/supportBoardFields.js": "게시판 스키마 필드",
    "server/lib/supportInquiryFields.js": "문의 스키마 필드",
    "server/lib/supportNewsFields.js": "소식 스키마 필드",
    "server/lib/supportNewsCommentFields.js": "소식 댓글 필드",
    "server/middleware/auth.js": "JWT 인증·requireRole",
    "server/middleware/supervisor.js": "슈퍼바이저 전용 가드",
    "server/scripts/copy-static.js": "루트 정적파일→server/public 복사",
    "server/scripts/ensure-pdf-font.js": "PDF 한글 폰트 준비",
    "server/scripts/ensure-vendor-new.js": "신규업체 컬렉션 초기화",
    "server/scripts/ensure-vendor-prospects.js": "예비업체 컬렉션 초기화",
    "server/scripts/test-mongo.js": "MongoDB 연결 테스트",
    "server/fonts/README.md": "PDF 한글 폰트 설치 안내",
    "server/assets/ak-seal.png": "AK 인감(서버 PDF용)",
    "server/assets/douzone-seal.png": "더존 인감(서버 PDF용)",
    "docs/ARCHITECTURE.md": "프로젝트 아키텍처 문서",
    "docs/thejohn-user-manual.docx": "사용자 매뉴얼(생성물)",
    "docs/thejohn-system-technical.docx": "기술 구조 문서(생성물)",
    "docs/thejohn-system-structure-management.docx": "시스템 구조 설명서(생성물)",
    "docs/thejohn-system-structure-management.pptx": "시스템 구조 요약(생성물)",
    "docs/thejohn-server-infrastructure.docx": "서버·인프라 문서(생성물)",
    "deploy/DEPLOY-GABIA.md": "가비아 배포 가이드",
    "deploy/RENDER-FIX.md": "Render 배포 이슈 해결",
    "deploy/MONGODB-FIX.md": "MongoDB 설정 가이드",
    "deploy/LOGIN-PERMISSIONS.md": "로그인·권한 설정",
    "deploy/STAFF-ACCOUNTS.md": "직원 계정 설정",
    "deploy/VENDOR-REGISTRATION.md": "업체 등록 가이드",
    "deploy/install-on-server.sh": "서버 설치 스크립트",
    "deploy/nginx-thejohn.conf.example": "Nginx 설정 예시",
    "deploy/ecosystem.config.cjs": "PM2 ecosystem 설정",
    "scripts/generate-user-manual.py": "사용자 매뉴얼 Word 생성",
    "scripts/generate-technical-doc.py": "기술 구조 Word 생성",
    "scripts/generate-management-docs.py": "구조·관리 문서·PPT 생성",
    "scripts/generate-server-doc.py": "서버·인프라 Word 생성",
    "scripts/generate-file-inventory-doc.py": "프로그램 파일 목록 Word 생성",
    "docs/thejohn-file-inventory.docx": "프로그램 파일 목록(생성물)",
    "docs/thejohn-deploy-gabia.docx": "배포 가비아 DNS·Render(생성물)",
    "docs/thejohn-deploy-render.docx": "배포 Render 설정(생성물)",
    "docs/thejohn-deploy-mongodb.docx": "배포 MongoDB 연결(생성물)",
    "docs/thejohn-deploy-login-permissions.docx": "배포 로그인·권한(생성물)",
    "docs/thejohn-deploy-staff-accounts.docx": "배포 직원 계정(생성물)",
    "docs/thejohn-deploy-vendor-registration.docx": "배포 업체 등록(생성물)",
    "scripts/generate-deploy-docs.py": "배포 가이드 Word 생성",
    "scripts/md_to_docx.py": "Markdown→Word 변환 공통",
    "scripts/file_inventory_data.py": "파일 목록·용도 데이터",
    "scripts/patch-html-brand-boot.js": "HTML brand-boot 패치",
    "scripts/sync-footer-html.js": "푸터 HTML 동기화",
    "requirements-docs.txt": "문서 생성 Python 의존성",
    "package.json": "루트 패키지(시작·빌드 진입)",
    "render.yaml": "Render Blueprint 배포 설정",
    "manifest.json": "PWA 웹앱 매니페스트",
    ".env.example": "환경변수 예시 템플릿",
    "README.md": "프로젝트 개요·실행 안내",
    ".github/workflows/deploy-render.yml": "Render 자동 배포 CI",
    "img/logo.png": "사이트 기본 로고",
    "img/icon-192.png": "PWA 아이콘 192",
    "img/icon-512.png": "PWA 아이콘 512",
    "img/intro.mp4": "메인 인트로 영상",
    "img/hub_intro.mp4": "허브 인트로 영상",
    "img/intro-music.mp3": "인트로 배경음",
    "img/hub_music.mp3": "허브 배경음",
    "img/seals/thejohn.png": "더존 인감 이미지",
    "img/seals/ak20140516.png": "AK 인감 이미지",
    "img/seals/README.txt": "인감 이미지 안내",
}


def _norm(path: str) -> str:
    return path.replace("\\", "/")


def _purpose(rel: str, category: str) -> str:
    key = _norm(rel)
    if key in PURPOSE_MAP:
        return PURPOSE_MAP[key]
    name = os.path.basename(key)
    if name in PURPOSE_MAP:
        return PURPOSE_MAP[name]
    if key.startswith("img/company-intro/"):
        return "회사소개 슬라이드 이미지"
    if category == "프론트 페이지 (HTML)":
        return "웹 화면 페이지"
    if category == "프론트 스크립트 (JavaScript)":
        return "화면 동작 스크립트"
    if category == "프론트 스타일 (CSS)":
        return "화면 스타일"
    if category.startswith("API 서버"):
        return "서버 프로그램"
    if category == "이미지·미디어 (img/)":
        return "이미지·미디어 자산"
    if category == "문서 (docs/)":
        return "프로젝트 문서"
    if category == "배포 가이드 (deploy/)":
        return "배포·운영 안내"
    if category == "유틸 스크립트 (scripts/)":
        return "유지보수·문서 생성 스크립트"
    return "프로젝트 구성 파일"


def _category(rel: str) -> str | None:
    rel = _norm(rel)
    if rel.startswith("server/public/"):
        return None
    if rel.endswith(".html") and "/" not in rel:
        return "프론트 페이지 (HTML)"
    if rel.endswith(".js") and "/" not in rel:
        return "프론트 스크립트 (JavaScript)"
    if rel.endswith(".css") and "/" not in rel:
        return "프론트 스타일 (CSS)"
    if rel in ("server/index.js", "server/db.js", "server/package.json"):
        return "API 서버 — 진입점·DB"
    if rel.startswith("server/routes/") and rel.endswith(".js"):
        return "API 서버 — routes"
    if rel.startswith("server/lib/") and rel.endswith(".js"):
        return "API 서버 — lib"
    if rel.startswith("server/middleware/") and rel.endswith(".js"):
        return "API 서버 — middleware"
    if rel.startswith("server/scripts/") and rel.endswith(".js"):
        return "API 서버 — scripts"
    if rel.startswith("server/fonts/") or rel.startswith("server/assets/"):
        return "API 서버 — fonts·assets"
    if rel.startswith("img/"):
        return "이미지·미디어 (img/)"
    if rel.startswith("docs/"):
        return "문서 (docs/)"
    if rel.startswith("deploy/"):
        return "배포 가이드 (deploy/)"
    if rel.startswith("scripts/") or rel == "requirements-docs.txt":
        return "유틸 스크립트 (scripts/)"
    if rel in (
        "package.json",
        "render.yaml",
        "manifest.json",
        ".env.example",
        "README.md",
    ) or rel.startswith(".github/"):
        return "프로젝트 설정·기타"
    return None


def iter_files(root: str = ROOT):
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = sorted(
            d for d in dirnames if d not in SKIP_DIRS and not d.startswith(".vscode")
        )
        for name in sorted(filenames):
            if name in SKIP_FILES:
                continue
            _, ext = os.path.splitext(name)
            if ext in SKIP_SUFFIXES:
                continue
            if name.startswith(".") and name not in (".env.example",):
                continue
            full = os.path.join(dirpath, name)
            rel = _norm(os.path.relpath(full, root))
            cat = _category(rel)
            if cat:
                yield cat, rel


def build_file_sections(root: str = ROOT):
    buckets: dict[str, list[tuple[str, str]]] = {}
    order = [
        "프론트 페이지 (HTML)",
        "프론트 스크립트 (JavaScript)",
        "프론트 스타일 (CSS)",
        "API 서버 — 진입점·DB",
        "API 서버 — routes",
        "API 서버 — lib",
        "API 서버 — middleware",
        "API 서버 — scripts",
        "API 서버 — fonts·assets",
        "이미지·미디어 (img/)",
        "문서 (docs/)",
        "배포 가이드 (deploy/)",
        "유틸 스크립트 (scripts/)",
        "프로젝트 설정·기타",
    ]
    for cat, rel in iter_files(root):
        buckets.setdefault(cat, []).append((rel, _purpose(rel, cat)))
    sections = []
    for cat in order:
        items = buckets.get(cat)
        if items:
            sections.append((cat, sorted(items, key=lambda x: x[0].lower())))
    return sections


def _set_cell_nowrap(cell):
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(OxmlElement("w:noWrap"))


def _set_table_fixed_layout(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _seq_col_width_cm(font_pt=9):
    """표 제목 '순번' 글자 폭의 1.5배 (9pt 한글 기준 약 0.28cm/자)."""
    header = "순번"
    char_cm = 0.28 * (font_pt / 9.0)
    return round(len(header) * char_cm * 1.5, 2)


def add_file_inventory_table(doc, rows, font_pt=9):
    """순번(제목 1.5배)·파일명(넓음·한 줄)·용도 표."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    headers = ["순번", "파일명", "용도"]
    table_total_cm = 16.1
    purpose_cm = 5.0
    seq_cm = round(_seq_col_width_cm(font_pt) * 2, 2)
    file_cm = round(table_total_cm - seq_cm - purpose_cm, 2)
    col_widths = (Cm(seq_cm), Cm(file_cm), Cm(purpose_cm))

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    _set_table_fixed_layout(table)

    for col_idx, width in enumerate(col_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width

    def fill_cell(cell, col_idx, text, is_header=False):
        cell.text = str(text)
        for paragraph in cell.paragraphs:
            if col_idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(font_pt)
                if col_idx == 1 and not is_header:
                    run.font.name = "Consolas"
        if col_idx == 1:
            _set_cell_nowrap(cell)

    for col_idx, header in enumerate(headers):
        fill_cell(table.rows[0].cells[col_idx], col_idx, header, is_header=True)

    for row_idx, row in enumerate(rows):
        for col_idx, val in enumerate(row):
            fill_cell(table.rows[row_idx + 1].cells[col_idx], col_idx, val)

    doc.add_paragraph()
